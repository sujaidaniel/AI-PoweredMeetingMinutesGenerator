from docx import Document

def save_to_word_document(content, file_name):
    # Create a new Word document
    document = Document()
    
    # Iterate over the dictionary and add content to the document
    for section, text in content.items():
        # Format section titles by replacing underscores with spaces and capitalizing words
        formatted_title = ' '.join(word.capitalize() for word in section.split('_'))
        document.add_heading(formatted_title, level=1)
        document.add_paragraph(text)
        # Add an empty paragraph for spacing between sections
        document.add_paragraph()
    
    # Save the document to the specified file name
    document.save(file_name)
